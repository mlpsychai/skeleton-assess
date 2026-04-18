"""
Report Generator

Generates formatted .docx reports from psychometric score calculations.
Uses data-driven category loop instead of per-category methods.
Embeds ECharts-rendered PNGs via chart_renderer + Playwright.
All instrument-specific values are read from instrument_config.json.
"""

import re
import tempfile
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from pathlib import Path
from typing import Dict, Any, Optional
from datetime import datetime

from .instrument_config import load_instrument_config, get_validity_category, get_substantive_categories, is_elevated


class ReportGenerator:
    """Generates formatted .docx score reports."""

    def __init__(self, instrument_config: Optional[Dict[str, Any]] = None):
        """
        Initialize the report generator.

        Args:
            instrument_config: Instrument configuration dict.
                               Loaded from instrument_config.json if not provided.
        """
        if instrument_config is None:
            instrument_config = load_instrument_config()

        self.config = instrument_config
        self.instrument_name = self.config['instrument_name']
        self.num_items = self.config['num_items']
        self.disclaimer_text = self.config.get('disclaimer_text', '')

        # Formatting from config
        fmt = self.config.get('formatting', {})
        self.docx_font = fmt.get('docx_font', 'Calibri')
        self.docx_table_style = fmt.get('docx_table_style', 'Light Grid Accent 1')
        status_colors = fmt.get('status_colors', {})
        self.color_valid = self._parse_rgb(status_colors.get('valid', '#008000'))
        self.color_warning = self._parse_rgb(status_colors.get('warning', '#FFA500'))
        self.color_invalid = self._parse_rgb(status_colors.get('invalid', '#FF0000'))

        # Validation threshold derived from config
        self.valid_completion_threshold = 1.0 - self.config['max_missing_threshold']

        # Categories from config
        self.substantive_categories = get_substantive_categories(self.config)
        val_cat = get_validity_category(self.config)
        self.validity_scales = val_cat['scales'] if val_cat else []

        # Chart renderer (optional — graceful degradation)
        self._chart_renderer = None
        try:
            from .chart_renderer import ChartRenderer
            self._chart_renderer = ChartRenderer(instrument_config=self.config)
        except Exception:
            pass

    def _apply_apa_table_style(self, table):
        """Apply APA 7th edition table formatting: horizontal rules only."""
        # Remove all existing borders
        tbl = table._tbl
        tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')

        # Set table-level borders: no vertical, horizontal only at top/bottom
        borders_xml = f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>
                <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>
                <w:left w:val="none" w:sz="0" w:space="0" w:color="000000"/>
                <w:right w:val="none" w:sz="0" w:space="0" w:color="000000"/>
                <w:insideH w:val="none" w:sz="0" w:space="0" w:color="000000"/>
                <w:insideV w:val="none" w:sz="0" w:space="0" w:color="000000"/>
            </w:tblBorders>
        '''
        tbl_pr.append(parse_xml(borders_xml))

        # Add bottom border to header row (row 0)
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                cell_borders = parse_xml(f'''
                    <w:tcBorders {nsdecls("w")}>
                        <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>
                    </w:tcBorders>
                ''')
                tc_pr.append(cell_borders)
                # Bold header text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

        # Remove all cell shading
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFFFFF" w:val="clear"/>')
                tc_pr.append(shading)

    def generate_report(self, calculation_results: Dict[str, Any],
                       validation_report: Dict[str, Any],
                       output_path: str,
                       narratives: Optional[Dict[str, str]] = None) -> str:
        """
        Generate a formatted .docx report.

        Args:
            calculation_results: Results from ScoreCalculator
            validation_report: Results from ScoreValidator
            output_path: Path for output .docx file
            narratives: Optional dict of narrative strings keyed by category

        Returns:
            Path to generated report
        """
        doc = Document()
        is_interpretive = narratives and any(narratives.values())

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = self.docx_font
        font.size = Pt(11)

        # Title
        report_type = "Interpretive Report" if is_interpretive else "Score Report"
        title = doc.add_heading(f'{self.instrument_name} {report_type}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Test Information Section
        self._add_test_info(doc, calculation_results, validation_report)

        # Validity Assessment Section
        self._add_validity_assessment(doc, validation_report)

        # Validity Scales Table
        self._add_scale_section(doc, calculation_results, "Validity Scales", self.validity_scales)

        # Data-driven loop over substantive categories
        for category in self.substantive_categories:
            self._add_scale_section(doc, calculation_results, category['title'], category['scales'])
            # Add per-category narrative after the table
            if narratives and category.get('key') in narratives:
                self._add_narrative(doc, narratives[category['key']])

        # Profile Integration
        if narratives and 'integration' in narratives:
            self._add_narrative_section(doc, 'Profile Integration', narratives['integration'])

        # Treatment Recommendations
        if narratives and 'treatment' in narratives:
            self._add_narrative_section(doc, 'Treatment Recommendations', narratives['treatment'])

        # Summary Section
        self._add_summary(doc, calculation_results)

        # Narrative Summary
        if narratives and 'summary' in narratives:
            self._add_narrative(doc, narratives['summary'])

        # Profile Graphs (ECharts via Playwright)
        self._add_profile_graphs(doc, calculation_results)

        # Disclaimer
        self._add_disclaimer(doc)

        # Appendix A — Scale Name Abbreviations
        self._add_appendix_a(doc, calculation_results)

        # Save document
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

        return str(output_path)

    def _add_test_info(self, doc, calc_results, val_report):
        """Add test information section."""
        doc.add_heading('Test Information', 1)

        table = doc.add_table(rows=4, cols=2)
        self._apply_apa_table_style(table)

        cells = table.rows[0].cells
        cells[0].text = 'Test ID'
        cells[1].text = calc_results['test_id']

        cells = table.rows[1].cells
        cells[0].text = 'Test Date'
        cells[1].text = calc_results['test_date']

        cells = table.rows[2].cells
        cells[0].text = 'Examinee ID'
        cells[1].text = calc_results['examinee_id']

        cells = table.rows[3].cells
        cells[0].text = 'Completion Rate'
        completion = val_report['completion_rate']
        items_completed = self.num_items - val_report['missing_count']
        cells[1].text = f"{completion:.1%} ({items_completed}/{self.num_items} items)"

        doc.add_paragraph()

    def _add_validity_assessment(self, doc, val_report):
        """Add validity assessment section."""
        doc.add_heading('Protocol Validity', 1)

        status = val_report['is_valid']
        completion = val_report['completion_rate']

        p = doc.add_paragraph()
        p.add_run('Status: ').bold = True

        if status and completion >= self.valid_completion_threshold:
            run = p.add_run('VALID')
            run.font.color.rgb = self.color_valid
        elif status:
            run = p.add_run('VALID with concerns')
            run.font.color.rgb = self.color_warning
        else:
            run = p.add_run('INVALID')
            run.font.color.rgb = self.color_invalid

        if val_report['warnings']:
            doc.add_paragraph('Warnings:', style='List Bullet')
            for warning in val_report['warnings']:
                doc.add_paragraph(warning, style='List Bullet 2')

        doc.add_paragraph()

    def _add_scale_section(self, doc, calc_results, section_title, scale_list):
        """Add a scale section with heading and table — generic for any category."""
        doc.add_heading(section_title, 1)

        scale_scores = calc_results['scale_scores']
        available_scales = [s for s in scale_list if s in scale_scores]

        if not available_scales:
            doc.add_paragraph('No scales available in this category.')
            doc.add_paragraph()
            return

        self._add_scale_table(doc, scale_scores, available_scales)

    def _add_scale_table(self, doc, scale_scores, scale_list):
        """Helper to add a scale table."""
        available_scales = [s for s in scale_list if s in scale_scores]

        if not available_scales:
            doc.add_paragraph('No scales available in this category.')
            doc.add_paragraph()
            return

        table = doc.add_table(rows=len(available_scales) + 1, cols=5)
        self._apply_apa_table_style(table)

        # Header row
        header = table.rows[0].cells
        header[0].text = 'Scale'
        header[1].text = 'Raw Score'
        header[2].text = self.config.get('score_label', 'T-Score')
        header[3].text = 'Items Scored'
        header[4].text = 'Interpretation'

        # Data rows
        for idx, scale_abbr in enumerate(available_scales, 1):
            scale_info = scale_scores[scale_abbr]
            cells = table.rows[idx].cells
            cells[0].text = f"{scale_abbr} - {scale_info['scale_name']}"
            cells[1].text = str(scale_info['raw_score'])
            t_display = scale_info.get('t_score_display', '')
            cells[2].text = t_display if t_display else 'N/A'
            cells[3].text = f"{scale_info['items_scored']}/{scale_info['total_items']}"
            cells[4].text = scale_info['interpretive_range']

            # Highlight elevated scores
            if is_elevated(self.config, scale_info['interpretive_range']):
                for cell in cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

        doc.add_paragraph()

    def _add_summary(self, doc, calc_results):
        """Add summary section."""
        doc.add_heading('Summary', 1)

        summary = calc_results['summary']

        doc.add_paragraph(f"Total scales: {summary['total_scales']}")
        doc.add_paragraph(f"Elevated scales: {summary['elevated_scales_count']}")

        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('Report generated: ').italic = True
        p.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S')).italic = True

        doc.add_paragraph()

    def _clean_narrative_text(self, text: str) -> str:
        """Clean markdown artifacts from LLM-generated narrative text."""
        lines = text.strip().split('\n')
        cleaned = []
        for line in lines:
            stripped = line.strip()
            # Skip markdown headings
            if re.match(r'^#{1,3}\s', stripped):
                continue
            # Strip bold markers
            stripped = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)
            # Strip bullet markers
            stripped = re.sub(r'^\s*-\s+', '', stripped)
            cleaned.append(stripped)
        return '\n'.join(cleaned)

    def _add_narrative(self, doc, narrative_text: str):
        """Add narrative paragraphs to the document."""
        cleaned = self._clean_narrative_text(narrative_text)
        paragraphs = cleaned.strip().split('\n\n')
        for para_text in paragraphs:
            text = para_text.strip().replace('\n', ' ')
            if text:
                p = doc.add_paragraph(text)
                p.paragraph_format.space_after = Pt(6)
        doc.add_paragraph()

    def _add_narrative_section(self, doc, heading: str, narrative_text: str):
        """Add a headed narrative section (integration, treatment, etc.)."""
        doc.add_heading(heading, 1)
        self._add_narrative(doc, narrative_text)

    def _add_profile_graphs(self, doc, calc_results):
        """Add profile graphs using ECharts PNGs rendered via Playwright."""
        doc.add_heading('Profile Visualization', 1)

        if self._chart_renderer is None:
            doc.add_paragraph(
                'Charts unavailable — install chart_renderer dependencies '
                '(playwright) for graph embedding.'
            )
            return

        scale_scores = calc_results['scale_scores']

        with tempfile.TemporaryDirectory() as tmpdir:
            png_files = self._chart_renderer.render_all_to_png(scale_scores, tmpdir)

            if not png_files:
                doc.add_paragraph(
                    'Graph rendering unavailable — Playwright may not be installed. '
                    'Use the HTML report for interactive charts.'
                )
                return

            # Combined profile
            if 'combined' in png_files:
                doc.add_picture(png_files['combined'], width=Inches(6.5))
                p = doc.add_paragraph()
                p.add_run(
                    'Visual representation of score elevations across all scales.'
                ).italic = True
                doc.add_paragraph()

            # Domain graphs
            domain_keys = [k for k in png_files if k.startswith('domain_')]
            if domain_keys:
                doc.add_page_break()
                doc.add_heading('Detailed Domain Graphs', 1)

                p = doc.add_paragraph()
                p.add_run(
                    'The following graphs show scales organized by clinical dysfunction domains.'
                ).italic = True
                doc.add_paragraph()

                for key in domain_keys:
                    domain_name = key.replace('domain_', '').replace('_', ' ').title()
                    doc.add_heading(domain_name, 2)
                    doc.add_picture(png_files[key], width=Inches(6.5))
                    doc.add_paragraph()

    def _add_disclaimer(self, doc):
        """Add disclaimer section."""
        doc.add_page_break()
        doc.add_heading('Disclaimer', 1)

        p = doc.add_paragraph(self.disclaimer_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _add_appendix_a(self, doc, calc_results):
        """Add Appendix A — Scale Name Abbreviations as two side-by-side tables."""
        doc.add_page_break()
        doc.add_heading('Appendix A', 1)

        p = doc.add_paragraph()
        run = p.add_run('Scale Name Abbreviations')
        run.italic = True

        scale_scores = calc_results['scale_scores']
        all_abbrevs = list(scale_scores.keys())
        mid = (len(all_abbrevs) + 1) // 2
        left = all_abbrevs[:mid]
        right = all_abbrevs[mid:]

        num_rows = max(len(left), len(right)) + 1
        table = doc.add_table(rows=num_rows, cols=4)
        self._apply_apa_table_style(table)

        # Header row
        header = table.rows[0].cells
        header[0].text = 'Abbreviation'
        header[1].text = 'Scale'
        header[2].text = 'Abbreviation'
        header[3].text = 'Scale'

        for idx, abbr in enumerate(left, 1):
            cells = table.rows[idx].cells
            cells[0].text = abbr
            cells[1].text = scale_scores[abbr]['scale_name']

        for idx, abbr in enumerate(right, 1):
            cells = table.rows[idx].cells
            cells[2].text = abbr
            cells[3].text = scale_scores[abbr]['scale_name']

    @staticmethod
    def _parse_rgb(hex_str: str) -> RGBColor:
        """Convert a hex color string like '#008000' to RGBColor."""
        hex_str = hex_str.lstrip('#')
        r, g, b = int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16)
        return RGBColor(r, g, b)
