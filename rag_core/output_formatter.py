"""
Output Formatter for RAG-generated content
Handles .docx generation and AI detection scoring
"""

import re
import math
from typing import Dict, Optional, Tuple
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class OutputFormatter:
    """Format RAG outputs with AI detection scores and .docx generation"""

    def __init__(self, default_format: str = "docx"):
        """
        Initialize output formatter

        Args:
            default_format: Default output format ("docx" or "txt")
        """
        self.default_format = default_format

    def calculate_ai_score(self, text: str) -> Dict[str, any]:
        """
        Calculate AI-likelihood score using heuristic analysis

        Analyzes text for:
        - Perplexity (word choice predictability)
        - Burstiness (sentence length variation)
        - Vocabulary diversity
        - Stylistic patterns common in AI text

        Args:
            text: Text to analyze

        Returns:
            Dict with ai_score (0-100), confidence, and breakdown
        """
        if not text or len(text.strip()) < 100:
            return {
                "ai_score": 0,
                "confidence": "low",
                "note": "Text too short for reliable analysis",
                "breakdown": {}
            }

        # Clean text
        text = text.strip()
        sentences = [s.strip() for s in re.split(r'[.!?]+', text) if s.strip()]
        words = text.lower().split()

        if len(sentences) < 3 or len(words) < 50:
            return {
                "ai_score": 0,
                "confidence": "low",
                "note": "Insufficient text for analysis",
                "breakdown": {}
            }

        # Calculate metrics
        metrics = {}

        # 1. Burstiness Score (0-100, lower = more AI-like)
        # AI text tends to have uniform sentence lengths
        sentence_lengths = [len(s.split()) for s in sentences]
        avg_length = sum(sentence_lengths) / len(sentence_lengths)
        variance = sum((l - avg_length) ** 2 for l in sentence_lengths) / len(sentence_lengths)
        std_dev = math.sqrt(variance)
        coefficient_of_variation = (std_dev / avg_length) if avg_length > 0 else 0

        # Low variation (< 0.3) suggests AI
        burstiness_score = min(100, coefficient_of_variation * 250)
        metrics['burstiness'] = 100 - burstiness_score  # Invert: high score = AI-like

        # 2. Vocabulary Diversity (0-100, lower = more AI-like)
        unique_words = len(set(words))
        total_words = len(words)
        diversity_ratio = unique_words / total_words if total_words > 0 else 0

        # AI tends to have moderate diversity (0.4-0.6)
        if 0.4 <= diversity_ratio <= 0.6:
            diversity_score = 70  # Suspicious range
        elif diversity_ratio < 0.4:
            diversity_score = 50  # Very repetitive (could be AI or poor writing)
        else:
            diversity_score = 30  # High diversity (more human-like)
        metrics['vocabulary_repetition'] = diversity_score

        # 3. Transition Word Density
        transitions = [
            'however', 'moreover', 'furthermore', 'additionally', 'consequently',
            'therefore', 'nevertheless', 'thus', 'hence', 'accordingly',
            'in conclusion', 'to summarize', 'in summary', 'specifically'
        ]
        transition_count = sum(1 for word in words if word in transitions)
        transition_density = (transition_count / total_words) * 1000 if total_words > 0 else 0

        # AI often overuses transitions (> 8 per 1000 words)
        transition_score = min(100, transition_density * 10)
        metrics['transition_overuse'] = transition_score

        # 4. Passive Voice Density
        passive_indicators = ['was', 'were', 'been', 'being', 'is', 'are']
        passive_count = sum(1 for word in words if word in passive_indicators)
        passive_density = (passive_count / total_words) * 100 if total_words > 0 else 0

        # AI tends toward moderate passive voice (3-7%)
        if 3 <= passive_density <= 7:
            passive_score = 60
        else:
            passive_score = 30
        metrics['passive_voice'] = passive_score

        # 5. Sentence Starter Patterns
        # AI often starts sentences with similar patterns
        starters = [s.split()[0].lower() for s in sentences if s.split()]
        unique_starters = len(set(starters))
        starter_diversity = unique_starters / len(starters) if starters else 1

        # Low starter diversity suggests AI
        starter_score = 100 - (starter_diversity * 100)
        metrics['sentence_pattern'] = starter_score

        # Calculate weighted AI score
        ai_score = (
            metrics['burstiness'] * 0.30 +
            metrics['vocabulary_repetition'] * 0.25 +
            metrics['transition_overuse'] * 0.20 +
            metrics['passive_voice'] * 0.15 +
            metrics['sentence_pattern'] * 0.10
        )

        # Determine confidence based on text length and consistency
        if total_words < 150:
            confidence = "low"
        elif total_words < 400:
            confidence = "medium"
        else:
            confidence = "high"

        # Determine category
        if ai_score < 30:
            category = "Likely Human"
            color = "green"
        elif ai_score < 60:
            category = "Uncertain/Mixed"
            color = "yellow"
        else:
            category = "Likely AI-Generated"
            color = "red"

        return {
            "ai_score": round(ai_score, 1),
            "confidence": confidence,
            "category": category,
            "color": color,
            "breakdown": {
                "uniformity_score": round(metrics['burstiness'], 1),
                "vocabulary_repetition": round(metrics['vocabulary_repetition'], 1),
                "transition_overuse": round(metrics['transition_overuse'], 1),
                "passive_voice": round(metrics['passive_voice'], 1),
                "sentence_patterns": round(metrics['sentence_pattern'], 1)
            },
            "text_stats": {
                "total_words": total_words,
                "unique_words": unique_words,
                "sentences": len(sentences),
                "avg_sentence_length": round(avg_length, 1)
            }
        }

    def create_docx(
        self,
        content: str,
        output_path: Path,
        title: Optional[str] = None,
        metadata: Optional[Dict] = None,
        include_ai_score: bool = True
    ) -> Path:
        """
        Create a formatted .docx file with AI detection score

        Args:
            content: Main text content
            output_path: Path for output file
            title: Document title
            metadata: Additional metadata to include
            include_ai_score: Whether to include AI detection score

        Returns:
            Path to created document
        """
        # Ensure .docx extension
        output_path = Path(output_path)
        if output_path.suffix.lower() != '.docx':
            output_path = output_path.with_suffix('.docx')

        # Create document
        doc = Document()

        # Calculate AI score if requested
        if include_ai_score:
            ai_analysis = self.calculate_ai_score(content)
        else:
            ai_analysis = None

        # Add AI Detection Score Header (if enabled)
        if ai_analysis:
            self._add_ai_score_header(doc, ai_analysis)

        # Add title if provided
        if title:
            title_para = doc.add_heading(title, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata if provided
        if metadata:
            self._add_metadata_section(doc, metadata)

        # Add main content
        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            # Check if it's a heading (all caps, short, or starts with #)
            if (para_text.isupper() and len(para_text) < 100) or para_text.startswith('#'):
                heading_text = para_text.lstrip('#').strip()
                doc.add_heading(heading_text, level=2)
            else:
                para = doc.add_paragraph(para_text)
                para_format = para.paragraph_format
                para_format.space_after = Pt(10)
                para_format.line_spacing = 2.0  # Double spacing

        # Add generation timestamp footer
        doc.add_page_break()
        footer_para = doc.add_paragraph()
        footer_run = footer_para.add_run(
            f"\n---\nGenerated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            f"Output formatted by RAG Core v0.1.0"
        )
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)

        # Save document
        doc.save(str(output_path))
        return output_path

    def _add_ai_score_header(self, doc: Document, ai_analysis: Dict) -> None:
        """Add AI detection score header to document"""
        # Add header box
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Title
        title_run = header.add_run("AI DETECTION ANALYSIS\n")
        title_run.bold = True
        title_run.font.size = Pt(14)

        # Score
        score_text = f"AI Likelihood Score: {ai_analysis['ai_score']}% "
        score_run = header.add_run(score_text)
        score_run.font.size = Pt(16)
        score_run.bold = True

        # Set color based on score
        if ai_analysis['color'] == 'green':
            score_run.font.color.rgb = RGBColor(0, 128, 0)
        elif ai_analysis['color'] == 'yellow':
            score_run.font.color.rgb = RGBColor(255, 140, 0)
        else:
            score_run.font.color.rgb = RGBColor(255, 0, 0)

        # Category
        category_run = header.add_run(f"({ai_analysis['category']})\n")
        category_run.font.size = Pt(12)
        category_run.italic = True

        # Confidence
        conf_run = header.add_run(f"Confidence: {ai_analysis['confidence'].title()}")
        conf_run.font.size = Pt(10)

        # Add breakdown
        breakdown_para = doc.add_paragraph()
        breakdown_para.add_run("Score Breakdown:\n").bold = True

        for metric, value in ai_analysis['breakdown'].items():
            metric_name = metric.replace('_', ' ').title()
            breakdown_para.add_run(f"  • {metric_name}: {value}%\n")

        # Add stats
        stats_para = doc.add_paragraph()
        stats_para.add_run("Text Statistics:\n").bold = True
        stats = ai_analysis['text_stats']
        stats_para.add_run(
            f"  • Total Words: {stats['total_words']}\n"
            f"  • Unique Words: {stats['unique_words']}\n"
            f"  • Sentences: {stats['sentences']}\n"
            f"  • Avg Sentence Length: {stats['avg_sentence_length']} words\n"
        )

        # Add disclaimer
        disclaimer = doc.add_paragraph()
        disclaimer_run = disclaimer.add_run(
            "\n⚠️ DISCLAIMER: This score is generated using heuristic analysis and should not be used "
            "as definitive proof of AI authorship. It is intended as a preliminary screening tool only.\n"
        )
        disclaimer_run.font.size = Pt(9)
        disclaimer_run.italic = True
        disclaimer_run.font.color.rgb = RGBColor(128, 128, 128)

        # Add separator
        doc.add_paragraph("─" * 80)
        doc.add_paragraph()  # Spacing

    def _add_metadata_section(self, doc: Document, metadata: Dict) -> None:
        """Add metadata section to document"""
        meta_para = doc.add_paragraph()
        meta_para.add_run("Document Information\n").bold = True

        for key, value in metadata.items():
            if value:
                key_formatted = key.replace('_', ' ').title()
                meta_para.add_run(f"{key_formatted}: {value}\n")

        doc.add_paragraph()  # Spacing

    def save_output(
        self,
        content: str,
        output_path: Path,
        format: Optional[str] = None,
        title: Optional[str] = None,
        metadata: Optional[Dict] = None,
        include_ai_score: bool = True
    ) -> Path:
        """
        Save output in specified format

        Args:
            content: Text content to save
            output_path: Path for output file
            format: Output format ("docx" or "txt"), defaults to self.default_format
            title: Document title
            metadata: Additional metadata
            include_ai_score: Whether to include AI detection score

        Returns:
            Path to saved file
        """
        format = format or self.default_format
        output_path = Path(output_path)

        if format.lower() == "docx":
            return self.create_docx(
                content=content,
                output_path=output_path,
                title=title,
                metadata=metadata,
                include_ai_score=include_ai_score
            )
        else:
            # Save as text with AI score header
            if include_ai_score:
                ai_analysis = self.calculate_ai_score(content)
                header = self._format_ai_score_text(ai_analysis)
                full_content = header + "\n\n" + "="*80 + "\n\n" + content
            else:
                full_content = content

            # Ensure .txt extension
            if output_path.suffix.lower() not in ['.txt', '.md']:
                output_path = output_path.with_suffix('.txt')

            output_path.write_text(full_content, encoding='utf-8')
            return output_path

    def _format_ai_score_text(self, ai_analysis: Dict) -> str:
        """Format AI score for text output"""
        lines = [
            "="*80,
            "AI DETECTION ANALYSIS".center(80),
            "="*80,
            f"\nAI Likelihood Score: {ai_analysis['ai_score']}% ({ai_analysis['category']})",
            f"Confidence: {ai_analysis['confidence'].title()}\n",
            "\nScore Breakdown:"
        ]

        for metric, value in ai_analysis['breakdown'].items():
            metric_name = metric.replace('_', ' ').title()
            lines.append(f"  • {metric_name}: {value}%")

        lines.append("\nText Statistics:")
        stats = ai_analysis['text_stats']
        lines.append(f"  • Total Words: {stats['total_words']}")
        lines.append(f"  • Unique Words: {stats['unique_words']}")
        lines.append(f"  • Sentences: {stats['sentences']}")
        lines.append(f"  • Avg Sentence Length: {stats['avg_sentence_length']} words")

        lines.append(
            "\n⚠️ DISCLAIMER: This score is generated using heuristic analysis and should "
            "not be used\nas definitive proof of AI authorship. It is intended as a preliminary "
            "screening tool only."
        )
        lines.append("="*80)

        return "\n".join(lines)
